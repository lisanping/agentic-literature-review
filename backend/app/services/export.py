"""Export service — generate Markdown and Word (.docx) outputs."""

import io
from pathlib import Path

import structlog

logger = structlog.stdlib.get_logger()


def export_markdown(
    content: str,
    references: list[dict] | None = None,
    title: str | None = None,
) -> str:
    """Build a complete Markdown document.

    Args:
        content: Main body text in Markdown format.
        references: List of reference dicts (each with at least "formatted").
        title: Optional document title.

    Returns:
        Complete Markdown text with title and references appended.
    """
    parts: list[str] = []
    if title:
        parts.append(f"# {title}\n")
    parts.append(content)

    if references:
        parts.append("\n---\n\n## References\n")
        for i, ref in enumerate(references, 1):
            formatted = ref.get("formatted", ref.get("title", ""))
            parts.append(f"[{i}] {formatted}\n")

    return "\n".join(parts)


def export_word(
    content: str,
    references: list[dict] | None = None,
    title: str | None = None,
) -> bytes:
    """Export content to a Word (.docx) document.

    Args:
        content: Main body text (Markdown-ish, will be rendered as paragraphs).
        references: List of reference dicts.
        title: Optional document title.

    Returns:
        Bytes of the generated .docx file.
    """
    from docx import Document
    from docx.shared import Pt

    doc = Document()

    if title:
        doc.add_heading(title, level=0)

    # Simple Markdown → Word conversion (headings + paragraphs)
    for line in content.split("\n"):
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("#### "):
            doc.add_heading(stripped[5:], level=4)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("- "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        else:
            doc.add_paragraph(stripped)

    if references:
        doc.add_heading("References", level=1)
        for i, ref in enumerate(references, 1):
            formatted = ref.get("formatted", ref.get("title", ""))
            p = doc.add_paragraph(f"[{i}] {formatted}")
            p.style.font.size = Pt(10)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def export_bibtex(references: list[dict]) -> str:
    """Export references as a BibTeX file."""
    from app.parsers.citation_formatter import CitationInfo, to_bibtex

    entries = []
    for ref in references:
        info = CitationInfo(
            title=ref.get("title", ""),
            authors=ref.get("authors", ["Unknown"]),
            year=ref.get("year"),
            venue=ref.get("venue"),
            doi=ref.get("doi"),
            url=ref.get("url"),
        )
        entries.append(to_bibtex(info))
    return "\n\n".join(entries)


def export_ris(references: list[dict]) -> str:
    """Export references as a RIS file."""
    from app.parsers.citation_formatter import CitationInfo, to_ris

    entries = []
    for ref in references:
        info = CitationInfo(
            title=ref.get("title", ""),
            authors=ref.get("authors", ["Unknown"]),
            year=ref.get("year"),
            venue=ref.get("venue"),
            doi=ref.get("doi"),
            url=ref.get("url"),
        )
        entries.append(to_ris(info))
    return "\n\n".join(entries)
